
from docx import Document
import csv

import io
import fitz  # PyMuPDF library
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from .forms import mf
from pdf2docx import Converter
import re
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from nltk.stem import PorterStemmer, WordNetLemmatizer
from pypdf import PdfReader
import pytesseract
from PIL import Image
from gensim.models.doc2vec import Doc2Vec, TaggedDocument
from numpy.linalg import norm
import numpy as np
import os
import pandas as pd



# Download necessary NLTK resources
nltk.download('punkt')  # For tokenization
nltk.download('stopwords')  # For stop words
nltk.download('wordnet')
# from pyabr.abbreviations import expand
from django.core.files.storage import FileSystemStorage
from django.core.exceptions import ObjectDoesNotExist




# calculating matching score for single resume
def similarity( jd_text,resume_text,model_name='cv_job_maching.model'):

    # Clean the text
    cleaned_resume = resume_text
    cleaned_jd = jd_text

    # Model evaluation
    model_name = 'cv_job_maching.model'
    model = Doc2Vec.load(model_name)
    v1 = model.infer_vector(cleaned_resume.split())
    v2 = model.infer_vector(cleaned_jd.split())
    similarity = 100 * (np.dot(np.array(v1), np.array(v2))) / (norm(np.array(v1)) * norm(np.array(v2)))

    return similarity


# calculating matching score for all resumes and display scores in descending order
def final_output(jd_path, pdf_folder_path ):
    jd_file = jd_path

    # List all files in the folder
    all_files = os.listdir(pdf_folder_path)

    # Filter to get only PDF files
    pdf_files = [filename for filename in all_files if filename.endswith(".pdf")]

    results = {}
    # Loop through each PDF file in the folder
    for pdf_file in pdf_files:
        # Get the full path to the PDF file
        pdf_path = os.path.join(pdf_folder_path, pdf_file)

        score = similarity(pdf_path, jd_file, model_name='cv_job_maching.model')

        results[pdf_file] = score

    # store all filename and along with similarity score
    final_result = pd.DataFrame(list(results.items()), columns=['Filename', 'Similarity Score'])

    # Sort the DataFrame by 'Score' in descending order
    df_sorted = final_result.sort_values(by='Similarity Score', ascending=False)

    return df_sorted
# Create your views here.
def pdf_to_docx(pdf_file):
    # Create Converter object
    cv = Converter(pdf_file)
    s = str(pdf_file)
    s1 = s.split(".")
    docx_file = s1[0] + ".docx"
    s = cv.convert(docx_file, start=0, end=None)
    cv.close()
    return docx_file


# extract text from pdf
def extract_text_from_pdf(pdf_path):
    text = ""
    file = PdfReader(pdf_path)
    pdf_reader = len(file.pages)
    for page_num in range(pdf_reader):
        page = file.pages[page_num]
        text += page.extract_text()
    return text


# save the dictionary data in csv file
def save_csv(data, csv_file):
    with open(csv_file, 'w', newline='') as f:
        w = csv.writer(f)
        w.writerow('heading', 'content')
        w.writerows(data)


# extract bold data from resume
def extract_bold_text(docx_file):
    document = Document(docx_file)
    bold_text_in_lines = []

    for paragraph in document.paragraphs:
        if paragraph.runs:
            line = ""
            font_size = paragraph.style.font.size.pt
            for run in paragraph.runs:
                if run.bold:
                    line += run.text
            if line:
                bold_text_in_lines.append(line.strip())

    return bold_text_in_lines


def expand(text):
    d = {}
    t = []
    with open(r"D:\oop\file_names.csv", 'r') as f:
        r = csv.reader(f)
        for i in r:
            a = i[0].split("-")
            t.append(a[0])
            b = a[1].replace("_", " ")
            a[1] = b
            d[a[0]] = a[1]

    for j in range(len(text)):
        s = text[j].upper()
        if s in t:
            r = d[s]
            text[j] = r

    y = ""
    for i in text:
        y += i + " "
    return y


# text processing
def preprocess_text(text, use_stemming=True, use_lemmatization=False):
    text = text.lower()

    # Remove punctuation and special characters
    text = re.sub(r'[^\w\s]', ' ', text)

    # Remove extra spaces
    text = re.sub(r'\s+', ' ', text).strip()

    # Tokenize the text into words
    words = word_tokenize(text)

    # Remove stop words
    stop_words = set(stopwords.words('english'))
    words = [word for word in words if word not in stop_words]

    # Apply stemming or lemmatization
    if use_stemming:
        stemmer = PorterStemmer()
        words = [stemmer.stem(word) for word in words]
    elif use_lemmatization:
        lemmatizer = WordNetLemmatizer()
        words = [lemmatizer.lemmatize(word) for word in words]

    # Join the words back into a cleaned text
    cleaned_text = ' '.join(words)
    s=cleaned_text.split()
    t=expand(s)
    return t


def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

#extract text from images
def extract_text_from_image(image_path):
    # Open the image file
    with Image.open(image_path) as img:
        # Use pytesseract to extract text
        text = pytesseract.image_to_string(img)
    return text
# def extract_text_from_txt(file):
#     with open(file,'r',encoding='utf-8') as f:
#         text=f.read()
#     return text
def extract_text_from_txt(file):
    try:
        if hasattr(file, 'read'):
            text = file.read().decode('utf-8')
        else:
            with open(file, 'r', encoding='utf-8') as f:
                text = f.read()
        return text
    except Exception as e:
        print(f"Error reading file {file}: {e}")
        return None

def resume_process(file):
    # Ensure the file name is extracted correctly for type checking
    file_name = file.name if hasattr(file, 'name') else str(file)

    if file_name.endswith(".pdf"):
        d = extract_text_from_pdf(file)
        d1 = preprocess_text(d)
        return d1
    elif file_name.endswith(".docx"):
        d = extract_text_from_docx(file)
        d1 = preprocess_text(d)
        return d1
    elif file_name.endswith('.txt'):
        text = extract_text_from_txt(file)
        if text is not None:
            processed_text = preprocess_text(text)
            return processed_text
        else:
            print(f"Failed to process file {file}")
            return None
    else:
        print(f"Unsupported file type: {file_name}")
        return None


def calculate_matching_percentage(job_description, resume):
    # Initialize a TF-IDF Vectorizer
    tfidf_vectorizer = TfidfVectorizer()

    # Fit and transform the job description and resume texts
    job_resume_matrix = tfidf_vectorizer.fit_transform([job_description, resume])

    # Calculate cosine similarity between the job description and resume vectors
    cosine_sim = cosine_similarity(job_resume_matrix)

    # The cosine similarity between the two documents is at position [0, 1] in the matrix
    similarity_score = cosine_sim[0, 1]


    return similarity_score
def simlarity(jd,r1):
    j=jd.split()
    r=r1.split()
    j1=set(j)
    r1=set(r)
    j=list(j1)
    r=list(r1)
    c=0
    l=len(j)
    for i in r:
        if i in j:
            c+=1
    p=(c/l)*100
    return p

def resumes_list(folder_path):
    # Folder path containing the PDFs
    pdf_folder_path = folder_path

    # List all files in the folder
    all_files = os.listdir(pdf_folder_path)

    # Filter to get only PDF files
    pdf_files = [filename for filename in all_files if filename.endswith((".pdf", ".docx",".txt"))]

    results = {}
    # Loop through each PDF file in the folder
    for pdf_file in pdf_files:
        # Get the full path to the PDF file
        pdf_path = os.path.join(pdf_folder_path, pdf_file)

    return pdf_files