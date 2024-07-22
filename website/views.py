from django.shortcuts import render,redirect,get_object_or_404

from django.contrib.auth import authenticate,login,logout
from django.http import HttpResponse
from .forms import mf
from django.contrib.auth.models import AnonymousUser, User
from django.views.decorators.cache import never_cache
from django.contrib.auth.decorators import login_required
from django.conf import settings
from django.contrib import messages
import glob
from django.core.files.storage import FileSystemStorage
from django.core.files.storage import default_storage
from django.contrib.auth.tokens import default_token_generator
import chardet
from .models import *
import os
import csv
 # collect all functions from functions.py in app website
# myapp/views.py
import spacy
from django.shortcuts import render

from django.http import JsonResponse

import joblib
from nltk.tokenize import word_tokenize
from sentence_transformers import SentenceTransformer, util

from docx import Document
import csv

import io
import fitz  # PyMuPDF library
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from .forms import mf
from pdf2docx import Converter
import re
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from nltk.stem import PorterStemmer, WordNetLemmatizer
from pypdf import PdfReader
import pytesseract
from PIL import Image
from gensim.models.doc2vec import Doc2Vec, TaggedDocument
from numpy.linalg import norm
import numpy as np
import os
import pandas as pd
from django.core.paginator import Paginator



# Download necessary NLTK resources
nltk.download('punkt')  # For tokenization
nltk.download('stopwords')  # For stop words
nltk.download('wordnet')
# from pyabr.abbreviations import expand
from django.core.files.storage import FileSystemStorage
from django.core.exceptions import ObjectDoesNotExist




# calculating matching score for single resume
def similarity( jd_text,resume_text,model_name='cv_job_maching.model'):

    # Clean the text
    cleaned_resume = resume_text
    cleaned_jd = jd_text

    # Model evaluation
    model_name = 'cv_job_maching.model'
    model = Doc2Vec.load(model_name)
    v1 = model.infer_vector(cleaned_resume.split())
    v2 = model.infer_vector(cleaned_jd.split())
    similarity = 100 * (np.dot(np.array(v1), np.array(v2))) / (norm(np.array(v1)) * norm(np.array(v2)))

    return similarity


# calculating matching score for all resumes and display scores in descending order
def final_output(jd_path, pdf_folder_path ):
    jd_file = jd_path

    # List all files in the folder
    all_files = os.listdir(pdf_folder_path)

    # Filter to get only PDF files
    pdf_files = [filename for filename in all_files if filename.endswith(".pdf")]

    results = {}
    # Loop through each PDF file in the folder
    for pdf_file in pdf_files:
        # Get the full path to the PDF file
        pdf_path = os.path.join(pdf_folder_path, pdf_file)

        score = similarity(pdf_path, jd_file, model_name='cv_job_maching.model')

        results[pdf_file] = score

    # store all filename and along with similarity score
    final_result = pd.DataFrame(list(results.items()), columns=['Filename', 'Similarity Score'])

    # Sort the DataFrame by 'Score' in descending order
    df_sorted = final_result.sort_values(by='Similarity Score', ascending=False)

    return df_sorted
# Create your views here.
def pdf_to_docx(pdf_file):
    # Create Converter object
    cv = Converter(pdf_file)
    s = str(pdf_file)
    s1 = s.split(".")
    docx_file = s1[0] + ".docx"
    s = cv.convert(docx_file, start=0, end=None)
    cv.close()
    return docx_file


# extract text from pdf
def extract_text_from_pdf(pdf_path):
    text = ""
    file = PdfReader(pdf_path)
    pdf_reader = len(file.pages)
    for page_num in range(pdf_reader):
        page = file.pages[page_num]
        text += page.extract_text()
    return text


# save the dictionary data in csv file
def save_csv(data, csv_file):
    with open(csv_file, 'w', newline='') as f:
        w = csv.writer(f)
        w.writerow('heading', 'content')
        w.writerows(data)


# extract bold data from resume
def extract_bold_text(docx_file):
    document = Document(docx_file)
    bold_text_in_lines = []

    for paragraph in document.paragraphs:
        if paragraph.runs:
            line = ""
            font_size = paragraph.style.font.size.pt
            for run in paragraph.runs:
                if run.bold:
                    line += run.text
            if line:
                bold_text_in_lines.append(line.strip())

    return bold_text_in_lines


def expand(text):
    d = {}
    t = []
    with open(r"D:\oop\file_names.csv", 'r') as f:
        r = csv.reader(f)
        for i in r:
            a = i[0].split("-")
            t.append(a[0])
            b = a[1].replace("_", " ")
            a[1] = b
            d[a[0]] = a[1]

    for j in range(len(text)):
        s = text[j].upper()
        if s in t:
            r = d[s]
            text[j] = r

    y = ""
    for i in text:
        y += i + " "
    return y


# text processing
def preprocess_text(text, use_stemming=True, use_lemmatization=False):
    text = text.lower()

    # Remove punctuation and special characters
    text = re.sub(r'[^\w\s]', ' ', text)

    # Remove extra spaces
    text = re.sub(r'\s+', ' ', text).strip()

    # Tokenize the text into words
    words = word_tokenize(text)

    # Remove stop words
    stop_words = set(stopwords.words('english'))
    words = [word for word in words if word not in stop_words]

    # Apply stemming or lemmatization
    if use_stemming:
        stemmer = PorterStemmer()
        words = [stemmer.stem(word) for word in words]
    elif use_lemmatization:
        lemmatizer = WordNetLemmatizer()
        words = [lemmatizer.lemmatize(word) for word in words]

    # Join the words back into a cleaned text
    cleaned_text = ' '.join(words)
    s=cleaned_text.split()
    t=expand(s)
    return t


def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

#extract text from images
def extract_text_from_image(image_path):
    # Open the image file
    with Image.open(image_path) as img:
        # Use pytesseract to extract text
        text = pytesseract.image_to_string(img)
    return text
# def extract_text_from_txt(file):
#     with open(file,'r',encoding='utf-8') as f:
#         text=f.read()
#     return text
def extract_text_from_txt(file):
    try:
        if hasattr(file, 'read'):
            text = file.read().decode('utf-8')
        else:
            with open(file, 'r', encoding='utf-8') as f:
                text = f.read()
        return text
    except Exception as e:
        print(f"Error reading file {file}: {e}")
        return None

def resume_process(file):
    # Ensure the file name is extracted correctly for type checking
    file_name = file.name if hasattr(file, 'name') else str(file)

    if file_name.endswith(".pdf"):
        d = extract_text_from_pdf(file)
        d1 = preprocess_text(d)
        return d1
    elif file_name.endswith(".docx"):
        d = extract_text_from_docx(file)
        d1 = preprocess_text(d)
        return d1
    elif file_name.endswith('.txt'):
        text = extract_text_from_txt(file)
        if text is not None:
            processed_text = preprocess_text(text)
            return processed_text
        else:
            print(f"Failed to process file {file}")
            return None
    else:
        print(f"Unsupported file type: {file_name}")
        return None


def calculate_matching_percentage(job_description, resume):
    # Initialize a TF-IDF Vectorizer
    tfidf_vectorizer = TfidfVectorizer()

    # Fit and transform the job description and resume texts
    job_resume_matrix = tfidf_vectorizer.fit_transform([job_description, resume])

    # Calculate cosine similarity between the job description and resume vectors
    cosine_sim = cosine_similarity(job_resume_matrix)

    # The cosine similarity between the two documents is at position [0, 1] in the matrix
    similarity_score = cosine_sim[0, 1]


    return similarity_score
def simlarity(jd,r1):
    j=jd.split()
    r=r1.split()
    j1=set(j)
    r1=set(r)
    j=list(j1)
    r=list(r1)
    c=0
    l=len(j)
    for i in r:
        if i in j:
            c+=1
    p=(c/l)*100
    return p

def resumes_list(folder_path):
    # Folder path containing the PDFs
    pdf_folder_path = folder_path

    # List all files in the folder
    all_files = os.listdir(pdf_folder_path)

    # Filter to get only PDF files
    pdf_files = [filename for filename in all_files if filename.endswith((".pdf", ".docx",".txt"))]

    results = {}
    # Loop through each PDF file in the folder
    for pdf_file in pdf_files:
        # Get the full path to the PDF file
        pdf_path = os.path.join(pdf_folder_path, pdf_file)

    return pdf_files
# Load pre-trained SBERT model
model = SentenceTransformer('paraphrase-MiniLM-L6-v2')

def compute_similarity_sbert(resume_text, job_description):
    # Encode sentences
    resume_embedding = model.encode(resume_text, convert_to_tensor=True)
    job_embedding = model.encode(job_description, convert_to_tensor=True)
    # Compute cosine similarity
    similarity = util.pytorch_cos_sim(resume_embedding, job_embedding)
    return similarity.item() * 100  # Convert to percentage

# Load the trained Doc2Vec model
# model_path = 'cv_job_maching.model'
# model = joblib.load(model_path)
#
# def compute_similarity_doc2vec(resume_text, job_description):
#     # Tokenize texts
#     resume_tokens = word_tokenize(resume_text.lower())
#     job_desc_tokens = word_tokenize(job_description.lower())
#     # Infer vectors
#     resume_vector = model.infer_vector(resume_tokens)
#     job_desc_vector = model.infer_vector(job_desc_tokens)
#     # Compute cosine similarity
#     similarity = model.dv.similarity_unseen_docs(resume_vector, job_desc_vector)
#     return similarity * 100  # Convert to percentage


ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx'}

def allowed_file(filename):
    ALLOWED_EXTENSIONS = {'txt', 'pdf', 'png', 'jpg', 'jpeg', 'gif'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

#main process of getting multi resume work percentage
def process_files(resume_directory, jd_filepath):
    results = []
    print(jd_filepath)
    jd= resume_process(jd_filepath)

    resume_files = glob.glob(os.path.join(resume_directory, '*'))
    for resume_file_path in resume_files:
        print(resume_file_path)
        r=resume_process(resume_file_path)
        score = similarity(r, jd)
        score=round(score,2)
        print(score)
        results.append({
            'filename': os.path.basename(resume_file_path),
            'score': score
        })

    results.sort(key=lambda x: x['score'], reverse=True)
    return results


# it will add the all upload files(which are through ui) in folder
@never_cache
@login_required(login_url='login')
def upload_file(request):
    results=[]
    selected=[]
    rejected=[]
    if request.method == 'POST':
        print("Request method is POST")

        # Handle resume files

        resume_files = request.FILES.getlist('resumes')
        print("+++++++===========resume files", resume_files)
        if not resume_files:
            print("No resume files found")
            return render(request, 'multi_resume.html', {'error': 'No resume file part'})
        jd_file = request.FILES.get('jd_file')
        if not jd_file:
            print("No JD file found")
            return render(request, 'multi_resume.html', {'error': 'No JD file part'})
        jd=resume_process(jd_file)
        for file in resume_files:
            print("+++++++===========",file)
            r = resume_process(file)
            score = compute_similarity_sbert(r, jd)
            score = round(score, 2)
            if score >60:
                selected.append({
                    'filename': file,
                    'score': score
                })
            else :
                rejected.append({
                    'filename': file,
                    'score': score
                })
            print(score)
            results.append({
                'filename': file,
                'score': score
            })

        # Handle JD file
    context={
        "results":results,
        "selected":selected,
        "rejected":rejected,
    }        # Process the uploaded files and get results

    return render(request, 'multi_resume.html',context)

#replace the old with new
def secure_filename(filename):
    return os.path.basename(filename).replace(' ', '_')



# Load spaCy model
nlp = spacy.load("en_core_web_md")

# calculate similarity using nlp
def compute_similarity(resume_text, job_description):
    # Convert texts to spaCy documents
    resume_doc = nlp(resume_text)
    job_desc_doc = nlp(job_description)
    # Compute similarity
    similarity = resume_doc.similarity(job_desc_doc)
    return similarity * 100  # Convert to percentage



# main functionality of matching
job_id=None
@never_cache
@login_required(login_url='login')
def matcher_view(request):
    d = Jd1.objects.all()
    print(d)
    m = Match_per.objects.all()
    p={}
    jd1=None
    resume=None
    matching_percentage = 0
    paginator = Paginator(m, 10)  # Show 10 employees per page

    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    if request.method == 'POST':
        job_id = request.POST.get('job_id')

        job_id=Jd1.objects.get(job_id=job_id)
        print("=============", job_id)
        job_description = get_object_or_404(Jd1, job_id=job_id)
        jd= job_description.description
        print("=============", jd)
        jd1=resume_process(jd)

        resume = request.FILES.get('resume')
        resume_content= resume_process(resume)
        print("===========+++++++++++",jd1)
        if jd1 and resume_content:
            print("++++++++++++++++++++++++++++++++")
            mp = compute_similarity_sbert(jd1, resume_content)
            matching_percentage = round(mp,2)
            print("++++++++++++============",matching_percentage)
            
        user = Match_per(job_id=job_id, resume=resume, mp=matching_percentage)
        p.update(jd=jd,resume=resume,matching_percentage= matching_percentage)
        if user:
            user.save()
            context={
                'descriptions': d,
                'match_per': m,
                'p':p,
                'page_obj': page_obj,
            }
            return render(request, 'result.html', context)
    context = {
        'descriptions': d,
        'match_per': m,
        'p':p,
        'page_obj': page_obj,
    }
    return render(request, 'result.html', context)






def  department_files(request, department):
    departments = ['MBSE', 'IT', 'CAE']
    department_dir = os.path.join(settings.CSV_FILES_DIR, department)
    csv_files = [f for f in os.listdir(department_dir) if f.endswith('.csv')]
    selected_file = request.GET.get('file')
    headers, data = [], []

    if selected_file:
        file_path = os.path.join(department_dir, selected_file)
        headers, data = read_csv_file(file_path)

    return render(request, 'home.html', {
        'department': department,
        'departments': departments,
        'csv_files': csv_files,
        'headers': headers,
        'data': data,
    })


def read_csv_file(file_path):
    data = []
    with open(file_path, newline='') as csvfile:
        reader = csv.reader(csvfile)
        headers = next(reader)  # Read the header row
        for row in reader:
            data.append(row)
    return headers, data

@never_cache
@login_required(login_url='login')
#show details from csv file in table formate
def csv_info(request):
    data = []
    headers=None
    if request.method == 'POST':
        dep = request.POST.get('dep')
        print("============",dep)
        file_name = f"{dep}.csv"
        print("==============",file_name)
        file_path=os.path.join(settings.CSV_FILES_DIR,file_name)
        if os.path.isfile(file_path):
            with open(file_path, newline='') as csvfile:
                reader = csv.reader(csvfile)
                headers = next(reader)  # Get headers
                # Find the indices of the required columns
                # indices = {header: index for index, header in enumerate(headers)}
                # required_headers = ['Experience',  'Positions', 'Status']
                # try:
                #     required_indices = [indices[header] for header in required_headers]
                # except KeyError as e:
                #     messages.error(request, f"Missing expected column: {e}")
                #     return render(request, 'home.html', {'headers': [], 'data': []})
                #
                # data = []
                # for row in reader:
                #     data.append([row[index] for index in required_indices])
                # headers = required_headers

                for row in reader:
                    data.append(row)
        else:
            messages.info(request, 'file is not their')

    context={
            'headers': headers,
            'data': data,
        }
    return render(request, 'home.html',context )


#delete code from history means previous matched percentage details from matcher_per table
@never_cache
@login_required(login_url='login')
def mp(request,id):
    p=Match_per.objects.get(id=id)
    p.delete()
    return redirect('matcher')

#add and show job description related
@never_cache
@login_required(login_url='login')
def addjd(request):
    descriptions=Jd1.objects.all()
    if request.method=='POST':
        job_id=request.POST.get('id')
        description=request.FILES.get('des')
        print(description)
        job_title=request.POST.get('title')
        user=Jd1(job_id=job_id,description=description,job_title=job_title)
        if user:
            user.save()
            messages.info(request, 'successfully added')
        else:
            messages.info(request,'fill all the fields')
    context={
        'descriptions':descriptions
    }
    return render(request,'jd.html',context)



#update of job description
@never_cache
@login_required(login_url='login')
def updatejd(request,id):
    description=get_object_or_404(Jd1,job_id=id)
    if request.method == 'POST':
        job_id=request.POST.get('id')
        description=request.FILES.get('description')
        title=request.POST.get('title')

        user=Jd1(job_id=job_id,description=description,job_title=title)
        user.save()
        return redirect('addjd')
    context={
        'descriptions':description
    }
    return render(request,'updatedescription.html',context)

# delete of job description
@never_cache
@login_required(login_url='login')
def deletejd(request,id):
    des=get_object_or_404(Jd1,job_id=id)
    p=Match_per.objects.all()
    if request.method == "POST":
        for i in p:
           if i==id:
               i.delete()
        des.delete()
        messages.info(request,"successfully deleted")
        return redirect('addjd')
    return render(request,"delete.html")

#search bar related code
@never_cache
@login_required(login_url='login')
def searchbar(request):
    if request.method=='GET':
        query=request.GET.get('query')
        if query:
            jd=Jd1.objects.filter(job_id__icontains=query)
            context={
                'descriptions':jd
            }
            return render(request,'searchbar.html',context)
    return render(request,'searchbar.html')
# employee page
def index(request):
    return render (request,"404.html")

@never_cache
@login_required(login_url='login')
def home(request):
    return render(request,'form.html')


#login page code
def logins(request):
    employ=Employee.objects.all()

    if request.method=='POST':
        username=request.POST.get('username')
        password=request.POST.get('password')
        user=authenticate(username=username,password=password)
        if user is not None:
            login(request, user)
            employee = Employee.objects.get(username=username)
            if employee.is_admin:
                # here we go for  users admin
                print("++++++++++++ true =================")

                return redirect('matcher')
            else:

                print("++++++++++++ false =================")

                return redirect('home')
            return redirect('index')

        else:
            messages.info(request,"username or password is wrong")
    return render(request,'login.html')

#logout related code
def logouts(request):
    logout(request)
    return redirect('index')
# def register(request):
#     if request.method == "POST":
#         email=request.POST.get('email')
#         name=request.POST.get('name')
#         password1=request.POST.get('password1')
#         password2=request.POST.get('password2')
#
#         if password1==password2:
#             if User.objects.filter(email=email).exists():
#                 messages.info(request,'user already exist')
#                 return redirect('register')
#             else:
#                 if User.objects.filter(username=name):
#                     messages.info(request, 'user already exist')
#                     return redirect('register')
#                 else:
#                     user = User.objects.create_user(email=email, username=name, password=password1)
#                     user.save()
#                     return redirect ('login')
#         else:
#             return redirect('register')
#
#     return render(request,'registration.html')


# @never_cache
# @login_required(login_url='login')
# def multi_matcher_view(request):
#     d = Jd1.objects.all()
#     matching_percentage = 0
#     if request.method == 'POST':
#         job_id = request.POST.get('job_id')
#         job_id=Jd1.objects.get(job_id=job_id)
#         job_description = get_object_or_404(Jd1, job_id=job_id)
#         jd = job_description.description
#         jd1=preprocess_text(jd)
#         r1=[]
#         resume1 = request.FILES.getlist('resume')
#         print(resume1)
#
#         for resume in resume1:
#             resume_content= resume_process(resume)
#             if jd and resume_content:
#                 print(jd)
#                 print(resume_content)
#                 mp = calculate_matching_percentage(jd1, resume_content)
#                 matching_percentage = mp*100
#             user = Match_per(job_id=job_id, resume=resume_content, mp=matching_percentage)
#             if user:
#                 user.save()
#             context = {
#             'matching_percentage': matching_percentage,
#             'jd':jd,
#             'resume':resume_content,
#                 }
#             return render(request, 'result.html', context)
#     context = {
#         'descriptions': d,
#
#     }
#     return render(request, 'multi_resume.html', context)